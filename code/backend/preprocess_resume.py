from typing import List, Optional
from datetime import datetime
from enum import Enum
import re
from pydantic import BaseModel, Field, validator, EmailStr
import instructor
from openai import OpenAI
from core.config import settings

# 配置 OpenAI 客户端
client = OpenAI(
    base_url=settings.ARK_BASE_URL,
    api_key=settings.ARK_API_KEY
)

client = instructor.from_openai(client)  # 添加 instructor 补丁

MODEL_NAME = settings.ARK_MODEL

# ============ 定义更精细的 Pydantic 模型 ============
class EducationLevel(str, Enum):
    BACHELOR = "本科"
    MASTER = "硕士"
    PHD = "博士"
    ASSOCIATE = "专科"
    OTHER = "其他"

class Education(BaseModel):
    school_name: str = Field(..., description="学校全称，与原件完全一致")
    degree: EducationLevel = Field(..., description="学位类型")
    major: str = Field(..., description="专业名称，与原件完全一致")
    start_date: str = Field(..., description="开始时间，YYYY-MM格式")
    end_date: str = Field(..., description="结束时间，YYYY-MM格式或'Present'")
    gpa: Optional[float] = Field(None, description="GPA成绩，如未提及则为空")
    
    @validator("start_date", "end_date")
    def validate_date_format(cls, v):
        if v != "Present" and not re.match(r"^\d{4}-(0[1-9]|1[0-2])$", v):
            raise ValueError(f"日期格式错误: {v}，应为 YYYY-MM 格式或 'Present'")
        return v

class WorkExperience(BaseModel):
    company_name: str = Field(..., description="公司全称")
    position: str = Field(..., description="职位名称")
    start_date: str = Field(..., description="开始时间，YYYY-MM格式")
    end_date: str = Field(..., description="结束时间，YYYY-MM格式或'Present'")
    location: Optional[str] = Field(None, description="工作地点")
    description: List[str] = Field(default_factory=list, description="工作职责和成就，每条为完整句子")
    skills_used: List[str] = Field(default_factory=list, description="在该工作中使用的技能")
    
    @validator("description")
    def validate_description_length(cls, v):
        if len(v) > 10:  # 限制最多10条描述
            return v[:10]
        return v

class Project(BaseModel):
    name: str = Field(..., description="项目名称")
    role: str = Field(..., description="担任角色")
    start_date: Optional[str] = Field(None, description="开始时间，YYYY-MM格式")
    end_date: Optional[str] = Field(None, description="结束时间，YYYY-MM格式或'Present'")
    description: List[str] = Field(default_factory=list, description="项目描述")
    technologies: List[str] = Field(default_factory=list, description="使用的技术栈")

class Skill(BaseModel):
    name: str = Field(..., description="技能名称")
    category: str = Field(..., description="技能类别，如'编程语言'、'框架'、'工具'等")
    proficiency: Optional[str] = Field(None, description="熟练程度：精通/熟练/了解")
    years_of_experience: Optional[float] = Field(None, description="使用年限")

class ContactInfo(BaseModel):
    phone: Optional[str] = Field(None, description="电话号码")
    email: Optional[EmailStr] = Field(None, description="邮箱地址")
    wechat: Optional[str] = Field(None, description="微信号")
    linkedin: Optional[str] = Field(None, description="LinkedIn链接")
    github: Optional[str] = Field(None, description="GitHub链接")

class ResumeData(BaseModel):
    """简历数据结构，支持完整的提取和验证"""
    
    # 基本信息
    name: str = Field(..., description="姓名")
    gender: Optional[str] = Field(None, description="性别")
    birth_date: Optional[str] = Field(None, description="出生日期，YYYY-MM-DD格式")
    
    # 联系信息
    contact: ContactInfo = Field(default_factory=ContactInfo)
    
    # 核心经历
    education: List[Education] = Field(default_factory=list, description="教育经历，按时间倒序排列")
    work_experience: List[WorkExperience] = Field(default_factory=list, description="工作经历，按时间倒序排列")
    projects: List[Project] = Field(default_factory=list, description="项目经历")
    
    # 技能
    skills: List[Skill] = Field(default_factory=list, description="技能列表")
    
    # 其他
    certifications: List[str] = Field(default_factory=list, description="证书列表")
    languages: List[str] = Field(default_factory=list, description="语言能力")
    self_introduction: Optional[str] = Field(None, description="自我评价/个人简介")
    
    # 验证
    @validator("work_experience", "education")
    def sort_by_date_desc(cls, v):
        """按时间倒序排列"""
        if not v:
            return v
        return sorted(v, key=lambda x: x.end_date if x.end_date != "Present" else "9999-12", reverse=True)
    
    def get_primary_skills(self, category: str = None) -> List[str]:
        """获取主要技能"""
        if category:
            return [skill.name for skill in self.skills if skill.category == category]
        return [skill.name for skill in self.skills if skill.proficiency in ["精通", "熟练"]]

# ============ 优化的简历提取函数 ============
def extract_resume_data(resume_text: str, max_retries: int = 3) -> ResumeData:
    """
    使用优化的提示词策略提取简历数据
    结合 CoT 推理和结构化约束
    """
    print(f"正在使用 {MODEL_NAME} 进行高精度简历解析...")
    
    # ===== 优化的系统提示词 =====
    system_prompt = """# Role
你是一款高性能的 AI 简历解析引擎，专门负责将非结构化的简历文本转化为高精度的结构化 JSON 数据。

# Extraction Rules (必须严格遵守)
1. **实体完整性**：提取工作经历时，必须保留完整的公司全称、职位名称和起止时间。
2. **时间标准化**：将所有日期转化为 YYYY-MM 格式（如 2023.05 -> 2023-05）。如果至今，请填入 "Present"。
3. **技能挖掘**：不仅提取明确列出的技能，还要从项目描述中推断核心技术栈（如提到 "使用 React 开发" -> 技能包含 "React"）。
4. **拒绝幻觉**：严禁编造简历中不存在的学校、公司或联系方式。如果某项信息缺失，请返回 null 或空列表。
5. **复杂文本处理**：如果简历中存在并列的多个教育背景或工作段落，请按时间倒序排列。
6. **上下文智能**：能区分"个人兴趣"与"专业技能"，区分"实习经历"与"正式工作"。

# 输出要求
请严格按照提供的 JSON Schema 输出结构，确保所有字段的类型和格式正确。"""

    # ===== 优化的用户提示词 =====
    user_prompt = f"""请分析下方简历文本，并提取结构化信息。

## 分析步骤（请按此逻辑思考）：
1. 首先，扫描整个文档，识别所有章节（个人信息、教育、工作、技能等）
2. 对于每个工作经历，仔细提取：公司全称、职位、时间、地点、职责描述
3. 从项目描述中挖掘技术关键词，并归类到技能
4. 验证联系方式的格式（邮箱、电话）
5. 确保所有时间已标准化为 YYYY-MM 格式
6. 最后，按 Schema 要求生成结构化 JSON

## 待处理简历文本：
--- 简历开始 ---
{resume_text}
--- 简历结束 ---

## 特别关注：
1. 从【项目描述】中提取具体技术实现，转化为结构化的技能点
2. 【教育背景】中的专业名称必须与原件完全一致
3. 如果存在多种时间格式（如"2023.5"、"2023年5月"、"May 2023"），统一转为"2023-05"
4. 对于仍在进行的经历，结束时间设为"Present"
5. 仔细区分正式工作和实习经历

请现在开始解析，确保输出完全符合 JSON Schema 要求。"""
    
    try:
        # 使用 instructor 的 structured output
        extraction_result = client.chat.completions.create(
            model=MODEL_NAME,
            response_model=ResumeData,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": user_prompt
                }
            ],
            temperature=0.1,  # 稍微提高一点创造性以处理复杂简历
            max_retries=max_retries,
            max_tokens=4000,  # 确保有足够tokens输出完整结果
        )
        
        print("✅ 简历解析完成！")
        return extraction_result
        
    except Exception as e:
        print(f"❌ 提取失败: {e}")
        # 尝试降级处理
        return fallback_extraction(resume_text)

# ============ 降级处理函数 ============
def fallback_extraction(resume_text: str) -> ResumeData:
    """当主解析失败时的降级方案"""
    print("⚠️ 使用降级解析...")
    
    # 简化的提示词
    simple_system = """你是一个简历解析器，从文本中提取基本信息。"""
    
    simple_user = f"""从以下简历中提取信息：
    
{resume_text}

请提取：姓名、邮箱、电话、教育背景、工作经历。如果某项信息缺失，请留空。"""
    
    try:
        result = client.chat.completions.create(
            model=MODEL_NAME,
            response_model=ResumeData,
            messages=[
                {"role": "system", "content": simple_system},
                {"role": "user", "content": simple_user}
            ],
            temperature=0,
        )
        return result
    except:
        # 返回空结构
        return ResumeData(name="未知")

# ============ 批处理和验证函数 ============
def batch_extract_resumes(resume_texts: List[str], batch_size: int = 5) -> List[ResumeData]:
    """批量处理简历"""
    results = []
    
    for i in range(0, len(resume_texts), batch_size):
        batch = resume_texts[i:i + batch_size]
        print(f"处理批次 {i//batch_size + 1}/{(len(resume_texts)-1)//batch_size + 1}")
        
        for resume in batch:
            try:
                result = extract_resume_data(resume)
                results.append(result)
            except Exception as e:
                print(f"处理失败: {e}")
                results.append(ResumeData(name="解析失败"))
    
    return results

def validate_resume_data(resume_data: ResumeData) -> dict:
    """验证提取的数据质量"""
    issues = []
    
    # 检查必填字段
    if not resume_data.name or resume_data.name == "未知":
        issues.append("缺少姓名")
    
    if not resume_data.contact.email:
        issues.append("缺少邮箱")
    
    if not resume_data.education:
        issues.append("缺少教育背景")
    
    if not resume_data.work_experience:
        issues.append("缺少工作经历")
    
    # 检查时间格式
    for edu in resume_data.education:
        if edu.end_date != "Present" and not re.match(r"^\d{4}-(0[1-9]|1[0-2])$", edu.end_date):
            issues.append(f"教育结束时间格式错误: {edu.end_date}")
    
    # 评分
    completeness_score = 100 - len(issues) * 10
    completeness_score = max(0, completeness_score)
    
    return {
        "has_issues": len(issues) > 0,
        "issues": issues,
        "completeness_score": completeness_score,
        "summary": {
            "name": resume_data.name,
            "email": resume_data.contact.email,
            "education_count": len(resume_data.education),
            "work_count": len(resume_data.work_experience),
            "skill_count": len(resume_data.skills)
        }
    }

# ============ 使用示例 ============
if __name__ == "__main__":
    import os
    import tempfile
    
    # 修改1：添加文件处理函数
    def extract_text_from_pdf(pdf_path: str) -> str:
        """从PDF文件中提取文本"""
        try:
            # 尝试使用pymupdf
            import fitz
            text = ""
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except ImportError:
            print("警告：pymupdf未安装，使用备用方案...")
            try:
                # 备用方案：使用pdfplumber
                import pdfplumber
                text = ""
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                print("错误：请安装pymupdf或pdfplumber")
                return ""
    
    def extract_text_from_word(docx_path: str) -> str:
        """从Word文件中提取文本"""
        try:
            import docx
            doc = docx.Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        except ImportError:
            print("错误：请安装python-docx")
            return ""
        except Exception as e:
            print(f"读取Word文件失败: {e}")
            return ""
    
    # 修改2：重构主函数，支持文件路径和文本输入
    def test_resume_parser(input_data: str, source_type: str = "text", file_path: str = None):
        """
        测试简历解析器
        :param input_data: 文本内容或文件路径
        :param source_type: "text" 或 "file"
        :param file_path: 如果source_type是file，指定文件路径
        """
        print("=" * 60)
        print("📄 简历解析测试开始")
        print("=" * 60)
        
        resume_text = input_data
        
        if source_type == "file":
            if not os.path.exists(input_data):
                print(f"❌ 文件不存在: {input_data}")
                return None
            
            file_ext = os.path.splitext(input_data)[1].lower()
            print(f"📁 处理文件: {input_data}")
            print(f"📄 文件类型: {file_ext}")
            
            if file_ext == '.pdf':
                resume_text = extract_text_from_pdf(input_data)
            elif file_ext in ['.docx', '.doc']:
                resume_text = extract_text_from_word(input_data)
            else:
                print(f"❌ 不支持的文件格式: {file_ext}")
                return None
            
            if not resume_text.strip():
                print("❌ 文件内容为空或提取失败")
                return None
        
        # 显示提取的原始文本（前500字符）
        print("\n📋 原始文本预览（前500字符）:")
        print("-" * 50)
        preview = resume_text[:500] + ("..." if len(resume_text) > 500 else "")
        print(preview)
        print("-" * 50)
        print(f"文本长度: {len(resume_text)} 字符")
        
        # 提取数据
        print(f"\n🤖 使用模型: {MODEL_NAME}")
        print("🔍 正在解析简历...")
        
        try:
            resume_data = extract_resume_data(resume_text)
            
            # 验证数据质量
            validation = validate_resume_data(resume_data)
            
            print("\n✅ 解析完成!")
            print("=" * 60)
            
            # 打印验证结果
            print(f"📊 数据质量评分: {validation['completeness_score']}/100")
            
            if validation['has_issues']:
                print(f"⚠️  发现的问题:")
                for issue in validation['issues']:
                    print(f"   - {issue}")
            else:
                print("🎉 数据完整，无问题")
            
            # 打印摘要信息
            print("\n📄 简历摘要:")
            print(f"   姓名: {resume_data.name}")
            print(f"   邮箱: {resume_data.contact.email}")
            print(f"   电话: {resume_data.contact.phone or '未填写'}")
            
            if resume_data.education:
                print(f"\n🎓 教育背景 ({len(resume_data.education)} 项):")
                for i, edu in enumerate(resume_data.education, 1):
                    print(f"   {i}. {edu.school_name} - {edu.major} ({edu.degree})")
                    print(f"      时间: {edu.start_date} 至 {edu.end_date}")
                    if edu.gpa:
                        print(f"      GPA: {edu.gpa}")
            
            if resume_data.work_experience:
                print(f"\n💼 工作经历 ({len(resume_data.work_experience)} 项):")
                for i, work in enumerate(resume_data.work_experience, 1):
                    print(f"   {i}. {work.company_name} - {work.position}")
                    print(f"      时间: {work.start_date} 至 {work.end_date}")
                    if work.location:
                        print(f"      地点: {work.location}")
                    if work.description:
                        print(f"      职责: {work.description[0][:50]}...")
            
            if resume_data.skills:
                print(f"\n🛠️  技能列表 ({len(resume_data.skills)} 项):")
                # 按类别分组
                skills_by_category = {}
                for skill in resume_data.skills:
                    if skill.category not in skills_by_category:
                        skills_by_category[skill.category] = []
                    skills_by_category[skill.category].append(
                        f"{skill.name}" + (f"({skill.proficiency})" if skill.proficiency else "")
                    )
                
                for category, skills in skills_by_category.items():
                    print(f"   {category}: {', '.join(skills[:5])}")
                    if len(skills) > 5:
                        print(f"      等共{len(skills)}项技能")
            
            if resume_data.projects:
                print(f"\n🚀 项目经历 ({len(resume_data.projects)} 项):")
                for i, project in enumerate(resume_data.projects[:3], 1):  # 只显示前3个
                    print(f"   {i}. {project.name} - {project.role}")
                    if project.technologies:
                        print(f"      技术栈: {', '.join(project.technologies[:5])}")
            
            # 打印详细数据结构
            print("\n🔍 详细数据结构:")
            print("-" * 50)
            
            # 1. 基本信息
            print("1. 基本信息:")
            print(f"   姓名: {resume_data.name}")
            print(f"   性别: {resume_data.gender or '未填写'}")
            print(f"   出生日期: {resume_data.birth_date or '未填写'}")
            
            # 2. 联系信息
            print("\n2. 联系信息:")
            contact = resume_data.contact
            print(f"   邮箱: {contact.email or '未填写'}")
            print(f"   电话: {contact.phone or '未填写'}")
            print(f"   微信: {contact.wechat or '未填写'}")
            print(f"   GitHub: {contact.github or '未填写'}")
            print(f"   LinkedIn: {contact.linkedin or '未填写'}")
            
            # 3. 核心技能
            print("\n3. 核心技能 (精通/熟练):")
            primary_skills = resume_data.get_primary_skills()
            if primary_skills:
                print(f"   {', '.join(primary_skills)}")
            else:
                print("   未识别到核心技能")
            
            # 4. 其他信息
            if resume_data.certifications:
                print(f"\n4. 证书 ({len(resume_data.certifications)} 项):")
                for cert in resume_data.certifications[:5]:
                    print(f"   - {cert}")
            
            if resume_data.languages:
                print(f"\n5. 语言能力:")
                print(f"   {', '.join(resume_data.languages)}")
            
            if resume_data.self_introduction:
                print(f"\n6. 自我评价 (前100字):")
                print(f"   {resume_data.self_introduction[:100]}...")
            
            print("=" * 60)
            return resume_data
            
        except Exception as e:
            print(f"❌ 解析过程中发生错误: {e}")
